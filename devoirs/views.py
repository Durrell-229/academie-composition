import os
import math
import json
import logging
from datetime import datetime, date
from io import BytesIO

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.db import transaction
from django.http import FileResponse, Http404, JsonResponse
from django.conf import settings
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from django.core.files.base import ContentFile

from .models import (
    Devoir, DevoirMatiere, DevoirComposition, DevoirReponseEleve,
    BulletinDevoir, BulletinDevoirLigne, Certificat
)
from core.models import Matiere, Classe
from accounts.models import User
from ai_engine.multi_ai import MultiAIService
from .forms import ClasseForm, HoraireForm

logger = logging.getLogger(__name__)

ADMIN_PHONE = '+2290197650817'


@login_required
def download_epreuve_file(request, pk):
    """Télécharge/voir une épreuve ou corrigé avec vérification."""
    from django.http import HttpResponse, HttpResponseForbidden
    
    dm = get_object_or_404(DevoirMatiere, pk=pk)
    file_type = request.GET.get('type', 'epreuve')
    
    # CORRECTION TYPE: Strict access control - NEVER allow students to see answer keys
    if file_type == 'corrige':
        if request.user.role not in ['admin', 'conseiller', 'professeur']:
            logger.warning(f"Élève {request.user.email} a tenté d'accéder au corrigé type de {dm}")
            messages.error(request, "Accès interdit : les corrigés types sont confidentiels.")
            return redirect('dashboard')
        # Even professors can only see their own corrigés
        if request.user.role == 'professeur' and dm.soumis_par != request.user:
            # Allow if admin has validated it
            if dm.statut != DevoirMatiere.StatutEP.VALIDE:
                messages.error(request, "Accès refusé au corrigé.")
                return redirect('dashboard')
    
    # ÉPREUVE: Allow if student has access to the devoir
    if file_type == 'epreuve':
        if request.user.role == 'eleve':
            # Check if devoir is in progress and for student's class
            if dm.devoir.statut != Devoir.Statut.EN_COURS:
                messages.error(request, "Cette épreuve n'est pas disponible.")
                return redirect('dashboard')
            user_classe = request.user.classe
            if user_classe and not dm.devoir.classes.filter(nom=user_classe).exists():
                messages.error(request, "Cette épreuve n'est pas pour votre classe.")
                return redirect('dashboard')
        elif request.user.role not in ['admin', 'conseiller']:
            if dm.soumis_par != request.user:
                messages.error(request, "Accès refusé.")
                return redirect('dashboard')
    
    # Déterminer quel fichier servir
    if file_type == 'corrige':
        file_field = dm.corrige_type_file
        filename_prefix = 'corrige'
    else:
        file_field = dm.epreuve_file
        filename_prefix = 'epreuve'
    
    if not file_field:
        messages.error(request, "Fichier non disponible.")
        return redirect('admin_validate_all')
    
    # Vérifier si le fichier existe physiquement
    try:
        if not file_field.storage.exists(file_field.name):
            logger.warning(f"Fichier {file_type} manquant pour DevoirMatiere {pk}")
            messages.error(request, f"Fichier {file_type} non disponible sur le serveur. Veuillez le soumettre à nouveau.")
            return redirect('admin_validate_all')
    except Exception as e:
        logger.error(f"Erreur vérification fichier {file_type}: {e}")
        messages.error(request, "Erreur lors de l'accès au fichier.")
        return redirect('admin_validate_all')
    
    # Servir le fichier
    try:
        response = FileResponse(
            file_field.open('rb'),
            content_type='application/pdf',
        )
        # Pour l'affichage dans le navigateur (pas as_attachment)
        if request.GET.get('view', '0') == '1':
            response['Content-Disposition'] = f'inline; filename="{filename_prefix}_{dm.matiere.nom}_{dm.devoir.titre[:20]}.pdf"'
        else:
            response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_{dm.matiere.nom}_{dm.devoir.titre[:20]}.pdf"'
        return response
    except Exception as e:
        logger.error(f"Erreur lecture fichier {file_type}: {e}")
        messages.error(request, "Erreur lors du téléchargement.")
        return redirect('admin_validate_all')


def _link_callback(uri, rel):
    """Convert media/static URLs to absolute file paths for xhtml2pdf."""
    import urllib.parse

    if not uri:
        return uri

    uri = urllib.parse.unquote(uri)

    if uri.startswith(settings.MEDIA_URL):
        rel_path = uri[len(settings.MEDIA_URL):]
        abs_path = os.path.join(str(settings.MEDIA_ROOT), rel_path)
        return os.path.normpath(abs_path)

    if uri.startswith(settings.STATIC_URL):
        rel_path = uri[len(settings.STATIC_URL):]
        abs_path = os.path.join(str(settings.STATIC_ROOT), rel_path)
        return os.path.normpath(abs_path)

    if uri.startswith('/media/'):
        rel_path = uri[7:]
        abs_path = os.path.join(str(settings.MEDIA_ROOT), rel_path)
        return os.path.normpath(abs_path)

    if uri.startswith('/static/'):
        rel_path = uri[8:]
        abs_path = os.path.join(str(settings.STATIC_ROOT), rel_path)
        return os.path.normpath(abs_path)

    return uri


def _get_ai_engine():
    """Get AI engine instance for correction."""
    try:
        return MultiAIService()
    except Exception as e:
        logger.error(f"Erreur initialisation IA: {e}")
        return None


def _extract_text_from_file(file_path):
    """Extract text from PDF or DOCX file for AI correction."""
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif ext == '.docx':
            import docx
            doc = docx.Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            # Try reading as text fallback
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
    except Exception as e:
        logger.error(f"Erreur extraction texte {ext}: {e}")
    return text


def _get_mention(moyenne):
    if moyenne >= 18:
        return 'Excellent — Tableau d\'Honneur'
    elif moyenne >= 16:
        return 'Très Bien'
    elif moyenne >= 14:
        return 'Bien'
    elif moyenne >= 12:
        return 'Assez Bien'
    elif moyenne >= 10:
        return 'Passable'
    else:
        return 'Insuffisant'


def _get_decision(moyenne):
    return 'Admis(e)' if moyenne >= 10 else 'Ajourné(e)'


# ═══════════════════════════════════════════
# ADMIN VIEWS
# ═══════════════════════════════════════════

@login_required
def devoir_list_view(request):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoirs = Devoir.objects.all().order_by('-date_debut')
    pending = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.SOUMIS).count()
    bulletins_en_attente = BulletinDevoir.objects.filter(statut=BulletinDevoir.StatutBulletin.EN_ATTENTE).count()

    return render(request, 'devoirs/admin_list.html', {
        'devoirs': devoirs,
        'pending_epreuves': pending,
        'bulletins_en_attente': bulletins_en_attente,
    })


@login_required
def devoir_create_view(request):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    if request.method == 'POST':
        titre = request.POST.get('titre', '').strip()
        description = request.POST.get('description', '').strip()
        date_debut = request.POST.get('date_debut', '')
        date_fin = request.POST.get('date_fin', '')
        instructions = request.POST.get('instructions', '').strip()
        coefficient_default = request.POST.get('coefficient_default', '1.00')
        classe_ids = request.POST.getlist('classes')
        matiere_ids = request.POST.getlist('matieres')

        if not all([titre, date_debut, date_fin]):
            messages.error(request, "Tous les champs sont requis.")
            return redirect('devoir_create')

        try:
            devoir = Devoir.objects.create(
                titre=titre,
                description=description,
                date_debut=date_debut,
                date_fin=date_fin,
                instructions=instructions,
                coefficient_default=coefficient_default,
                createur=request.user,
            )
            if classe_ids:
                devoir.classes.set(classe_ids)
            if matiere_ids:
                devoir.matieres.set(matiere_ids)

            messages.success(request, f"Devoir '{titre}' créé avec succès.")
            return redirect('devoir_detail', pk=devoir.pk)
        except Exception as e:
            messages.error(request, f"Erreur: {str(e)}")

    classes = Classe.objects.all().order_by('nom')
    matieres = Matiere.objects.all().order_by('nom')
    return render(request, 'devoirs/admin_create.html', {
        'classes': classes,
        'matieres': matieres,
    })


@login_required
def devoir_detail_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=pk)
    matieres_info = []
    for m in devoir.matieres.all():
        dm = DevoirMatiere.objects.filter(devoir=devoir, matiere=m).first()
        coeff = dm.coefficient if dm else devoir.coefficient_default
        matieres_info.append({
            'matiere': m,
            'epreuve': dm,
            'coefficient': coeff,
            'horaire': {
                'date': dm.date_epreuve if dm else None,
                'start': dm.heure_debut if dm else None,
                'end': dm.heure_fin if dm else None,
                'salle': dm.salle if dm else '',
            },
        })

    bulletins_en_attente = BulletinDevoir.objects.filter(devoir=devoir, statut=BulletinDevoir.StatutBulletin.EN_ATTENTE).count()
    bulletins_approuves = BulletinDevoir.objects.filter(devoir=devoir, statut=BulletinDevoir.StatutBulletin.APPROUVE).count()

    return render(request, 'devoirs/admin_detail.html', {
        'devoir': devoir,
        'matieres_info': matieres_info,
        'bulletins_en_attente': bulletins_en_attente,
        'bulletins_approuves': bulletins_approuves,
    })


@login_required
def devoir_publish_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    devoir = get_object_or_404(Devoir, pk=pk)
    devoir.statut = Devoir.Statut.PROGRAMME_PUBLIE
    devoir.save()
    
    # NOTIFICATION: Notify profs and CP that a new devoir is published
    from notifications.models import Notification
    profs = User.objects.filter(role=User.Role.PROFESSEUR, is_active=True)
    cps = User.objects.filter(role=User.Role.CONSEILLER, is_active=True)
    
    for user in profs | cps:
        Notification.objects.create(
            recipient=user,
            title=f"Nouveau devoir publié: {devoir.titre}",
            message=f"Le devoir '{devoir.titre}' a été publié. Veuillez soumettre les épreuves pour vos matières.",
            type='INSCRIPTION',
        )
    
    messages.success(request, "Programme publié. Profs et CP notifiés.")
    return redirect('devoir_detail', pk=pk)


@login_required
def devoir_start_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    devoir = get_object_or_404(Devoir, pk=pk)
    devoir.statut = Devoir.Statut.EN_COURS
    devoir.save()
    for classe in devoir.classes.all():
        eleves = User.objects.filter(role=User.Role.ELEVE, classe=classe.nom)
        for eleve in eleves:
            DevoirComposition.objects.get_or_create(
                devoir=devoir, eleve=eleve, defaults={'classe': classe}
            )
    messages.success(request, "Devoir lancé, élèves inscrits.")
    return redirect('devoir_detail', pk=pk)


@login_required
def devoir_end_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    devoir = get_object_or_404(Devoir, pk=pk)
    devoir.statut = Devoir.Statut.TERMINE
    devoir.save()
    messages.success(request, "Devoir terminé.")
    return redirect('devoir_detail', pk=pk)


@login_required
def devoir_matiere_validate_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    dm = get_object_or_404(DevoirMatiere, pk=pk)
    dm.statut = DevoirMatiere.StatutEP.VALIDE
    dm.valide_par = request.user
    dm.validated_at = timezone.now()
    dm.save()
    
    # NOTIFICATION: Notify students that epreuve is validated and available
    from notifications.models import Notification
    
    # Get class names (strings) from the devoir's classes
    classe_names = dm.devoir.classes.values_list('nom', flat=True)
    
    eleves = User.objects.filter(
        role=User.Role.ELEVE, 
        is_active=True,
        classe__in=classe_names
    ).distinct()
    
    for eleve in eleves:
        Notification.objects.create(
            recipient=eleve,
            title=f"Épreuve disponible: {dm.matiere.nom}",
            message=f"L'épreuve de {dm.matiere.nom} pour le devoir '{dm.devoir.titre}' est maintenant disponible. Vous pouvez composer.",
            type='INSCRIPTION',
        )
    
    messages.success(request, f"Épreuve '{dm.matiere.nom}' validée. Élèves notifiés.")
    return redirect('devoir_detail', pk=dm.devoir.pk)


@login_required
def devoir_matiere_reject_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    dm = get_object_or_404(DevoirMatiere, pk=pk)
    dm.statut = DevoirMatiere.StatutEP.REJETE
    dm.commentaire_admin = request.POST.get('commentaire', '')
    dm.save()
    messages.warning(request, f"Épreuve '{dm.matiere.nom}' rejetée.")
    return redirect('devoir_detail', pk=dm.devoir.pk)


@login_required
def admin_validate_all_view(request):
    """Vue globale pour valider/rejeter toutes les soumissions d'épreuves."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    statut = request.GET.get('statut', 'soumis')
    
    epreuves = DevoirMatiere.objects.select_related(
        'devoir', 'matiere', 'soumis_par', 'valide_par'
    ).order_by('-submitted_at')
    
    if statut == 'tous':
        pass
    elif statut == 'soumis':
        epreuves = epreuves.filter(statut=DevoirMatiere.StatutEP.SOUMIS)
    elif statut == 'valide':
        epreuves = epreuves.filter(statut=DevoirMatiere.StatutEP.VALIDE)
    elif statut == 'rejete':
        epreuves = epreuves.filter(statut=DevoirMatiere.StatutEP.REJETE)
    
    # Stats
    pending_count = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.SOUMIS).count()
    valides_count = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.VALIDE).count()
    rejetes_count = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.REJETE).count()
    total_count = DevoirMatiere.objects.count()
    
    return render(request, 'devoirs/admin_validate.html', {
        'epreuves': epreuves,
        'statut': statut,
        'pending_count': pending_count,
        'valides_count': valides_count,
        'rejetes_count': rejetes_count,
        'total_count': total_count,
    })


@login_required
def admin_upload_epreuve_view(request, pk):
    """Admin directly uploads epreuve + corrige for a matiere."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=pk)
    matiere_id = request.GET.get('matiere_id')
    matiere = get_object_or_404(Matiere, pk=matiere_id) if matiere_id else None

    if request.method == 'POST':
        matiere_id = request.POST.get('matiere_id')
        matiere = get_object_or_404(Matiere, pk=matiere_id)
        epreuve = request.FILES.get('epreuve')
        corrige = request.FILES.get('corrige')

        if not epreuve or not corrige:
            messages.error(request, "L'épreuve et le corrigé sont requis.")
            return redirect('admin_upload_epreuve', pk=pk, matiere_id=matiere_id)

        dm, created = DevoirMatiere.objects.update_or_create(
            devoir=devoir, matiere=matiere,
            defaults={
                'epreuve_file': epreuve,
                'corrige_type_file': corrige,
                'soumis_par': request.user,
                'statut': DevoirMatiere.StatutEP.SOUMIS,
            }
        )
        messages.success(request, f"Épreuve '{matiere.nom}' soumise. À valider.")
        return redirect('devoir_detail', pk=pk)

    return render(request, 'devoirs/admin_upload_epreuve.html', {
        'devoir': devoir,
        'matiere': matiere,
        'matieres': devoir.matieres.all(),
    })


@login_required
def devoir_set_coefficients_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=pk)

    if request.method == 'POST':
        coefficients = {}
        for m in devoir.matieres.all():
            coeff_str = request.POST.get(f'coeff_{m.id}', str(devoir.coefficient_default))
            try:
                coeff = float(coeff_str)
                if coeff > 0:
                    coefficients[str(m.id)] = coeff
            except ValueError:
                coefficients[str(m.id)] = float(devoir.coefficient_default)

        devoir.coefficients_par_matiere = coefficients
        devoir.save()
        messages.success(request, "Coefficients mis à jour avec succès.")
        return redirect('devoir_detail', pk=pk)

    return render(request, 'devoirs/admin_set_coefficients.html', {
        'devoir': devoir,
    })


@login_required
def devoir_bulletins_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=pk)
    statut_filter = request.GET.get('statut', 'en_attente')

    bulletins = BulletinDevoir.objects.filter(devoir=devoir).select_related('eleve', 'classe')
    if statut_filter == 'en_attente':
        bulletins = bulletins.filter(statut=BulletinDevoir.StatutBulletin.EN_ATTENTE)
    elif statut_filter == 'approuve':
        bulletins = bulletins.filter(statut=BulletinDevoir.StatutBulletin.APPROUVE)
    elif statut_filter == 'rejete':
        bulletins = bulletins.filter(statut=BulletinDevoir.StatutBulletin.REJETE)
    elif statut_filter == 'tous':
        pass

    bulletins = bulletins.order_by('-moyenne_generale')

    return render(request, 'devoirs/admin_bulletins.html', {
        'devoir': devoir,
        'bulletins': bulletins,
        'statut_filter': statut_filter,
    })


@login_required
def devoir_approuver_bulletin_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    bulletin = get_object_or_404(BulletinDevoir, pk=pk)
    if request.method == 'POST':
        # Admin can adjust grades before approval
        for ligne in bulletin.lignes.all():
            note_finale_str = request.POST.get(f'note_finale_{ligne.id}')
            if note_finale_str:
                try:
                    note_finale = float(note_finale_str)
                    note_finale = max(0, min(20, note_finale))
                    ligne.moyenne = note_finale
                    ligne.save()
                except ValueError:
                    pass

        # Recalculate average
        _recalculer_moyenne_bulletin(bulletin)

        bulletin.statut = BulletinDevoir.StatutBulletin.APPROUVE
        bulletin.approuve_par = request.user
        bulletin.approuve_at = timezone.now()
        bulletin.save()

        # Generate PDF
        _generer_bulletin_pdf(bulletin)

        messages.success(request, f"Bulletin de {bulletin.eleve.full_name} approuvé.")
    return redirect('devoir_bulletins', pk=bulletin.devoir.pk)


@login_required
def devoir_rejeter_bulletin_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    bulletin = get_object_or_404(BulletinDevoir, pk=pk)
    if request.method == 'POST':
        commentaire = request.POST.get('commentaire', '')
        bulletin.statut = BulletinDevoir.StatutBulletin.REJETE
        bulletin.appreciation_ia = commentaire
        bulletin.save()
        messages.warning(request, f"Bulletin de {bulletin.eleve.full_name} rejeté.")
    return redirect('devoir_bulletins', pk=bulletin.devoir.pk)


@login_required
def devoir_approuver_tous_view(request, pk):
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=pk)
    if request.method == 'POST':
        bulletins = BulletinDevoir.objects.filter(
            devoir=devoir, statut=BulletinDevoir.StatutBulletin.EN_ATTENTE
        )
        count = 0
        for bulletin in bulletins:
            bulletin.statut = BulletinDevoir.StatutBulletin.APPROUVE
            bulletin.approuve_par = request.user
            bulletin.approuve_at = timezone.now()
            bulletin.save()
            _generer_bulletin_pdf(bulletin)
            count += 1
        messages.success(request, f"{count} bulletin(s) approuvé(s).")
    return redirect('devoir_bulletins', pk=pk)


# ═══════════════════════════════════════════
# DELETE VIEWS (Admin only)
# ═══════════════════════════════════════════

@login_required
def delete_epreuve_view(request, pk):
    """Supprimer une épreuve (DevoirMatiere) - Admin uniquement."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    dm = get_object_or_404(DevoirMatiere, pk=pk)
    devoir_pk = dm.devoir.pk
    matiere_nom = dm.matiere.nom
    
    # Check if there are student copies
    copies_count = DevoirReponseEleve.objects.filter(devoir_matiere=dm).count()
    if copies_count > 0:
        messages.error(
            request,
            f"Impossible de supprimer: {copies_count} copie(s) d'élève(s) déjà soumise(s). "
            f"Supprimez d'abord les copies des élèves."
        )
        return redirect('admin_validate_all')
    
    # Check if bulletin exists
    bulletins_count = BulletinDevoir.objects.filter(devoir_matiere=dm).count()
    if bulletins_count > 0:
        messages.error(
            request,
            f"Impossible de supprimer: {bulletins_count} bulletin(s) déjà généré(s)."
        )
        return redirect('admin_validate_all')
    
    if request.method == 'POST':
        logger.info(f"Admin {request.user.email} supprime l'épreuve {matiere_nom} de {dm.devoir.titre}")
        dm.delete()
        messages.success(request, f"Épreuve de {matiere_nom} supprimée.")
    
    return redirect('admin_validate_all')


@login_required
def delete_devoir_view(request, pk):
    """Supprimer un devoir complet - Admin uniquement."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    devoir = get_object_or_404(Devoir, pk=pk)
    
    # Check if students have composed
    compositions_count = DevoirComposition.objects.filter(devoir=devoir).count()
    if compositions_count > 0:
        messages.error(
            request,
            f"Impossible de supprimer: {compositions_count} composition(s) d'élève(s) existe(nt). "
            f"Supprimez d'abord les compositions."
        )
        return redirect('devoir_list')
    
    # Check if bulletins exist
    bulletins_count = BulletinDevoir.objects.filter(devoir=devoir).count()
    if bulletins_count > 0:
        messages.error(
            request,
            f"Impossible de supprimer: {bulletins_count} bulletin(s) déjà généré(s)."
        )
        return redirect('devoir_list')
    
    if request.method == 'POST':
        logger.info(f"Admin {request.user.email} supprime le devoir {devoir.titre}")
        devoir_nom = devoir.titre
        devoir.delete()
        messages.success(request, f"Devoir '{devoir_nom}' supprimé.")
    
    return redirect('devoir_list')


# ═══════════════════════════════════════════
# PROF VIEWS
# ═══════════════════════════════════════════

@login_required
def prof_dashboard_view(request):
    """Prof dashboard — shows devoirs with epreuve submission status."""
    if request.user.role not in ('professeur', 'admin'):
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    # Get all active devoirs (brouillon, programme_publie, en_cours)
    devoirs = Devoir.objects.filter(
        statut__in=[Devoir.Statut.BROUILLON, Devoir.Statut.PROGRAMME_PUBLIE, Devoir.Statut.EN_COURS]
    ).prefetch_related('matieres').order_by('-date_debut')

    # Pre-fetch all DevoirMatiere records to avoid N+1 queries
    devoir_ids = [d.pk for d in devoirs]
    dm_qs = DevoirMatiere.objects.filter(devoir_id__in=devoir_ids).select_related('matiere')
    dm_map = {}
    for dm in dm_qs:
        dm_map[(dm.devoir_id, dm.matiere_id)] = dm

    # Build data with submission status — no extra DB hit per (devoir, matiere)
    devoirs_data = []
    for devoir in devoirs:
        for matiere in devoir.matieres.all():
            dm = dm_map.get((devoir.pk, matiere.pk))
            devoirs_data.append({
                'devoir': devoir,
                'matiere': matiere,
                'epreuve': dm,
            })

    return render(request, 'devoirs/prof_dashboard.html', {
        'devoirs_data': devoirs_data,
    })


@login_required
def devoir_submit_epreuve_view(request, devoir_id, matiere_id):
    if request.user.role not in ('professeur', 'admin'):
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    devoir = get_object_or_404(Devoir, pk=devoir_id)
    matiere = get_object_or_404(Matiere, pk=matiere_id)

    if request.method == 'POST':
        epreuve = request.FILES.get('epreuve')
        corrige = request.FILES.get('corrige')

        if not epreuve or not corrige:
            messages.error(request, "L'épreuve et le corrigé sont requis.")
            return redirect('devoir_submit_epreuve', devoir_id=devoir_id, matiere_id=matiere_id)

        DevoirMatiere.objects.update_or_create(
            devoir=devoir, matiere=matiere,
            defaults={
                'epreuve_file': epreuve,
                'corrige_type_file': corrige,
                'soumis_par': request.user,
                'statut': DevoirMatiere.StatutEP.SOUMIS,
            }
        )
        
        # NOTIFICATION: Notify admin that prof has submitted
        from notifications.models import Notification
        admins = User.objects.filter(role=User.Role.ADMIN, is_active=True)
        for admin in admins:
            Notification.objects.create(
                recipient=admin,
                title=f"Épreuve soumise: {matiere.nom}",
                message=f"Le professeur {request.user.full_name} a soumis l'épreuve de {matiere.nom} pour le devoir '{devoir.titre}'.",
                type='INSCRIPTION',
            )
        
        messages.success(request, "Épreuve soumise avec succès. Admin notifié.")
        return redirect('prof_dashboard')

    return render(request, 'devoirs/prof_submit.html', {
        'devoir': devoir,
        'matiere': matiere,
    })


# ═══════════════════════════════════════════
# ELEVE VIEWS
# ═══════════════════════════════════════════

@login_required
def eleve_programme_view(request):
    """Show all published/active devoirs for the student's class."""
    user_classe = request.user.classe
    
    logger.info(f"[Eleve Programme] User: {request.user.email}, classe: '{user_classe}'")
    
    if not user_classe:
        # Student has no class assigned - show all published devoirs with a warning
        logger.warning(f"[Eleve Programme] User {request.user.email} has no classe assigned")
        devoirs = Devoir.objects.filter(
            statut__in=[Devoir.Statut.PROGRAMME_PUBLIE, Devoir.Statut.EN_COURS],
        ).distinct().order_by('date_debut')
        messages.info(request, "Vous n'avez pas de classe assignée. Voici tous les devoirs publiés.")
    else:
        # Filter devoirs by student's class
        devoirs = Devoir.objects.filter(
            statut__in=[Devoir.Statut.PROGRAMME_PUBLIE, Devoir.Statut.EN_COURS],
            classes__nom=user_classe,
        ).distinct().order_by('date_debut')
        
        logger.info(f"[Eleve Programme] Found {devoirs.count()} devoirs for classe '{user_classe}'")

    return render(request, 'devoirs/eleve_programme.html', {'devoirs': devoirs})


@login_required
def eleve_devoir_calendar_view(request, devoir_id):
    """Vue calendrier détaillée pour un devoir avec horaires par matière."""
    devoir = get_object_or_404(Devoir, pk=devoir_id)
    
    # Vérifier que l'élève est dans les classes concernées
    user_classe = request.user.classe
    logger.info(f"[Eleve Calendar] User: {request.user.email}, classe: '{user_classe}', devoir classes: {list(devoir.classes.values_list('nom', flat=True))}")
    
    if user_classe and not devoir.classes.filter(nom=user_classe).exists():
        messages.error(request, "Ce devoir n'est pas pour votre classe.")
        return redirect('eleve_programme')
    
    # Récupérer les matières validées avec leurs horaires
    matieres_valides = DevoirMatiere.objects.filter(
        devoir=devoir, 
        statut=DevoirMatiere.StatutEP.VALIDE
    ).select_related('matiere').order_by('devoir__date_debut')
    
    # Construire le planning détaillé
    planning = []
    for dm in matieres_valides:
        horaire = devoir.horaires.get(dm.matiere.nom, {})
        planning.append({
            'matiere': dm.matiere,
            'devoir_matiere': dm,
            'date': horaire.get('date', str(devoir.date_debut)),
            'heure_demarrage': horaire.get('heure_demarrage', '08:00'),
            'heure_fin': horaire.get('heure_fin', '10:00'),
            'salle': horaire.get('salle', ''),
            'coeff': devoir.coefficients_par_matiere.get(str(dm.matiere.id), devoir.coefficient_default),
        })
    
    # Récupérer les réponses existantes de l'élève
    reponses = DevoirReponseEleve.objects.filter(
        devoir_matiere__devoir=devoir, 
        eleve=request.user
    )
    reponses_map = {r.devoir_matiere_id: r for r in reponses}
    
    # Composition de l'élève
    composition, _ = DevoirComposition.objects.get_or_create(
        devoir=devoir, 
        eleve=request.user,
        defaults={'classe': devoir.classes.filter(nom=request.user.classe).first()}
    )
    
    return render(request, 'devoirs/eleve_calendar.html', {
        'devoir': devoir,
        'planning': planning,
        'composition': composition,
        'reponses_map': reponses_map,
    })


@login_required
def eleve_compose_view(request, devoir_id):
    """Show the composition page for a specific devoir."""
    devoir = get_object_or_404(Devoir, pk=devoir_id)
    
    # Verify student's class matches
    user_classe = request.user.classe
    logger.info(f"[Eleve Compose] User: {request.user.email}, classe: '{user_classe}', devoir classes: {list(devoir.classes.values_list('nom', flat=True))}")
    
    if user_classe and not devoir.classes.filter(nom=user_classe).exists():
        messages.error(request, "Ce devoir n'est pas pour votre classe.")
        return redirect('eleve_programme')
    
    composition, _ = DevoirComposition.objects.get_or_create(
        devoir=devoir, eleve=request.user,
        defaults={'classe': devoir.classes.filter(nom=user_classe).first() or devoir.classes.first()}
    )

    if devoir.statut != Devoir.Statut.EN_COURS:
        messages.error(request, "Ce devoir n'est pas en cours.")
        return redirect('eleve_programme')

    matieres_valides = DevoirMatiere.objects.filter(
        devoir=devoir, statut=DevoirMatiere.StatutEP.VALIDE
    )

    # Get existing responses
    reponses = DevoirReponseEleve.objects.filter(
        devoir_matiere__devoir=devoir, eleve=request.user
    )
    reponses_map = {r.devoir_matiere_id: r for r in reponses}

    return render(request, 'devoirs/eleve_compose.html', {
        'devoir': devoir,
        'composition': composition,
        'matieres': matieres_valides,
        'reponses_map': reponses_map,
    })


@login_required
def eleve_upload_copie_view(request, devoir_matiere_id):
    """Upload a copy for a specific subject."""
    dm = get_object_or_404(DevoirMatiere, pk=devoir_matiere_id)

    if dm.devoir.statut != Devoir.Statut.EN_COURS:
        messages.error(request, "Ce devoir n'est pas en cours.")
        return redirect('eleve_programme')

    if request.user.role != 'eleve':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    reponse, created = DevoirReponseEleve.objects.get_or_create(
        devoir_matiere=dm, eleve=request.user,
        defaults={'statut': DevoirReponseEleve.StatutReponse.SOUMIS}
    )

    if request.method == 'POST':
        copie_file = request.FILES.get('copie_file')

        if not copie_file:
            messages.error(request, "Veuillez uploader votre copie.")
            return redirect('eleve_upload_copie', devoir_matiere_id=devoir_matiere_id)

        reponse.copie_file = copie_file
        reponse.statut = DevoirReponseEleve.StatutReponse.SOUMIS
        reponse.save()

        # Trigger IA auto-correction
        _corriger_copie_ia(reponse, dm)

        messages.success(request, "Copie soumise et en cours de correction par l'IA.")
        return redirect('eleve_compose', devoir_id=dm.devoir.pk)

    return render(request, 'devoirs/eleve_upload_copie.html', {
        'dm': dm,
        'reponse': reponse,
        'devoir': dm.devoir,
    })


@login_required
def eleve_submit_reponse_view(request, devoir_id):
    """Submit all copies for a devoir — triggers bulletin generation."""
    if request.method != 'POST':
        return redirect('eleve_programme')

    devoir = get_object_or_404(Devoir, pk=devoir_id)
    composition, _ = DevoirComposition.objects.get_or_create(
        devoir=devoir, eleve=request.user,
        defaults={'classe': devoir.classes.first()}
    )

    if devoir.statut != Devoir.Statut.EN_COURS:
        messages.error(request, "Ce devoir n'est pas en cours.")
        return redirect('eleve_programme')

    # Check if all copies are submitted and corrected
    matieres_valides = DevoirMatiere.objects.filter(
        devoir=devoir, statut=DevoirMatiere.StatutEP.VALIDE
    )
    reponses = DevoirReponseEleve.objects.filter(
        devoir_matiere__in=matieres_valides, eleve=request.user
    )

    total_submitted = reponses.filter(
        statut__in=[DevoirReponseEleve.StatutReponse.CORRIGE, DevoirReponseEleve.StatutReponse.APPROUVE]
    ).count()

    if total_submitted < matieres_valides.count():
        messages.warning(
            request,
            f"Vous avez soumis {total_submitted}/{matieres_valides.count()} matières. "
            "Soumettez toutes vos copies avant de finaliser."
        )
    else:
        # Generate bulletin for this student
        _generer_bulletin_pour_eleve(devoir, request.user, composition.classe)
        composition.statut = DevoirComposition.StatutComp.COMPOSE
        composition.composed_at = timezone.now()
        composition.save()
        messages.success(request, "Toutes vos copies ont été soumises. Bulletin en attente d'approbation.")

    return redirect('eleve_resultats')


@login_required
def eleve_resultats_view(request):
    compositions = DevoirComposition.objects.filter(
        eleve=request.user
    ).select_related('devoir').order_by('-devoir__date_debut')

    # Get bulletins for this student
    bulletins = BulletinDevoir.objects.filter(eleve=request.user).select_related('devoir').order_by('-created_at')

    certificats = Certificat.objects.filter(eleve=request.user).order_by('-created_at')

    return render(request, 'devoirs/eleve_resultats.html', {
        'compositions': compositions,
        'bulletins': bulletins,
        'certificats': certificats,
    })


@login_required
def eleve_download_bulletin_view(request, pk):
    """Student downloads their own bulletin — strict ownership check."""
    bulletin = get_object_or_404(BulletinDevoir, pk=pk)

    if bulletin.eleve != request.user:
        messages.error(request, "Accès refusé. Ce bulletin ne vous appartient pas.")
        return redirect('dashboard')

    if bulletin.statut != BulletinDevoir.StatutBulletin.APPROUVE:
        messages.error(request, "Ce bulletin n'est pas encore approuvé.")
        return redirect('eleve_resultats')

    if not bulletin.file_pdf:
        _generer_bulletin_pdf(bulletin)

    if not bulletin.file_pdf:
        raise Http404("Bulletin non disponible")

    return FileResponse(
        bulletin.file_pdf.open('rb'),
        content_type='application/pdf',
        as_attachment=True,
        filename=f"bulletin_{bulletin.eleve.full_name.replace(' ', '_')}_{bulletin.devoir.titre[:20]}.pdf"
    )


# ═══════════════════════════════════════════
# CERTIFICATE VIEWS
# ═══════════════════════════════════════════

@login_required
def devoir_certificates_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    devoir = get_object_or_404(Devoir, pk=pk)
    compositions = DevoirComposition.objects.filter(
        devoir=devoir, statut=DevoirComposition.StatutComp.COMPOSE
    ).order_by('rang')
    return render(request, 'devoirs/admin_certificates.html', {
        'devoir': devoir,
        'compositions': compositions,
    })


@login_required
def devoir_generate_certificates_view(request, pk):
    if request.user.role != 'admin':
        return redirect('dashboard')
    devoir = get_object_or_404(Devoir, pk=pk)

    bulletins_approuves = BulletinDevoir.objects.filter(
        devoir=devoir, statut=BulletinDevoir.StatutBulletin.APPROUVE
    )

    count = 0
    for bulletin in bulletins_approuves:
        cert, created = Certificat.objects.get_or_create(
            eleve=bulletin.eleve,
            devoir=devoir,
            type_certificat=Certificat.TypeCertificat.ADMISSION,
            defaults={
                'moyenne_obtenue': bulletin.moyenne_generale,
                'mention': _get_mention(bulletin.moyenne_generale),
            }
        )
        if created or not cert.file_pdf:
            _generer_certificat_devoir_pdf(cert, bulletin)
            count += 1

    messages.success(request, f"{count} certificat(s) généré(s).")
    return redirect('devoir_certificates', pk=pk)


@login_required
def certificat_download_view(request, pk):
    cert = get_object_or_404(Certificat, pk=pk)
    if request.user.role != 'admin' and cert.eleve != request.user:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')

    if not cert.file_pdf:
        # Try to find associated bulletin
        bulletin = BulletinDevoir.objects.filter(
            devoir=cert.devoir, eleve=cert.eleve
        ).first()
        if bulletin:
            _generer_certificat_devoir_pdf(cert, bulletin)
        else:
            _generate_certificat_pdf(cert)

    if not cert.file_pdf:
        raise Http404("Certificat non disponible")

    return FileResponse(
        cert.file_pdf.open('rb'),
        content_type='application/pdf',
        as_attachment=True,
        filename=f"certificat_{cert.numero_certificat}.pdf"
    )


# ═══════════════════════════════════════════
# IA CORRECTION SERVICE
# ═══════════════════════════════════════════

def _corriger_copie_ia(reponse, devoir_matiere):
    """Run IA auto-correction on a student's copy using the specific answer key."""
    try:
        logger.info(f"[IA Correction] Démarrage pour {reponse.eleve.full_name} - {devoir_matiere.matiere.nom}")
        logger.info(f"[IA Correction] Devoir: {devoir_matiere.devoir.titre}")
        
        engine = _get_ai_engine()
        if not engine:
            logger.warning("[IA Correction] IA non disponible, copie en attente de correction manuelle.")
            reponse.statut = DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION
            reponse.save()
            return

        # Extract text from the student's copy
        copie_text = ""
        if reponse.copie_file:
            copie_path = reponse.copie_file.path
            logger.info(f"[IA Correction] Copie fichier: {copie_path}")
            copie_text = _extract_text_from_file(copie_path)
            if not copie_text and reponse.copie_text:
                copie_text = reponse.copie_text
                logger.info(f"[IA Correction] Copie texte direct: {len(copie_text)} caractères")
        
        logger.info(f"[IA Correction] Copie extraite: {len(copie_text)} caractères")

        # Extract text from the answer key (corrigé type)
        corrige_text = ""
        corrige_doc_id = getattr(devoir_matiere, 'corrige_document_id', '')
        logger.info(f"[IA Correction] Corrigé document ID: {corrige_doc_id}")
        
        if devoir_matiere.corrige_type_file:
            corrige_path = devoir_matiere.corrige_type_file.path
            logger.info(f"[IA Correction] Corrigé fichier trouvé: {corrige_path}")
            corrige_text = _extract_text_from_file(corrige_path)
            logger.info(f"[IA Correction] Corrigé extrait: {len(corrige_text)} caractères")
        else:
            logger.error(f"[IA Correction] PAS DE CORRIGÉ TYPE pour {devoir_matiere.matiere.nom} - fichier manquant!")

        if not corrige_text:
            logger.error(f"[IA Correction] ÉCHEC: Pas de corrigé type extrait pour {devoir_matiere.matiere.nom}")
            reponse.statut = DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION
            reponse.save()
            return
        
        logger.info(f"[IA Correction] SUCCÈS: Corrigé type récupéré et extrait ({len(corrige_text)} caractères)")

        # Get coefficient for this subject
        coeff = devoir_matiere.devoir.coefficients_par_matiere.get(
            str(devoir_matiere.matiere.id),
            devoir_matiere.devoir.coefficient_default
        )

        copie_doc_id = getattr(reponse, 'copie_document_id', '')

        exam_info = {
            'titre': devoir_matiere.devoir.titre,
            'matiere': devoir_matiere.matiere.nom,
            'note_maximale': 20,
            'niveau': devoir_matiere.devoir.classes.first().nom if devoir_matiere.devoir.classes.exists() else 'Secondaire',
            'coefficient': float(coeff),
            'corrige_doc_id': corrige_doc_id,
            'copie_doc_ids': [copie_doc_id] if copie_doc_id else [],
            'session_id': str(reponse.id),
        }

        logger.info(f"[Devoirs IA] Envoi à IA: corrige={corrige_doc_id} copie={copie_doc_id} eleve={reponse.eleve.full_name}")
        logger.info(f"[Devoirs IA] Corrigé: {len(corrige_text)} caractères, Copie: {len(copie_text)} caractères")

        # Call AI correction
        result = engine.correct_copy(corrige_text, copie_text, exam_info)
        
        logger.info(f"[Devoirs IA] Réponse IA reçue: {type(result)}")

        if isinstance(result, dict) and 'note' in result:
            note = float(result.get('note', 0))
            note = max(0, min(20, note))

            reponse.note_ia = round(note, 2)
            reponse.note_finale = round(note, 2)
            reponse.appreciation_ia = result.get('appreciation', '')
            reponse.feedback_ia = result
            reponse.statut = DevoirReponseEleve.StatutReponse.CORRIGE
            reponse.corrige_at = timezone.now()
            reponse.save()

            logger.info(f"IA correction for {reponse.eleve.full_name} - {devoir_matiere.matiere.nom}: {note}/20")
        else:
            logger.error(f"Résultat IA invalide: {result}")
            reponse.statut = DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION
            reponse.save()

    except Exception as e:
        logger.error(f"Erreur correction IA pour {reponse.eleve.full_name}: {e}")
        reponse.statut = DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION
        reponse.save()


# ═══════════════════════════════════════════
# BULLETIN GENERATION SERVICE
# ═══════════════════════════════════════════

def _generer_bulletin_pour_eleve(devoir, eleve, classe):
    """Generate a BulletinDevoir for a student from their corrected copies."""
    try:
        # bulletins supprimé
        get_benin_coefficient = lambda *a, **k: 1
        serie = classe or (eleve.classe if eleve else '')
        
        with transaction.atomic():
            # Get all corrected responses for this student
            reponses = DevoirReponseEleve.objects.filter(
                devoir_matiere__devoir=devoir,
                eleve=eleve,
                statut__in=[DevoirReponseEleve.StatutReponse.CORRIGE, DevoirReponseEleve.StatutReponse.APPROUVE]
            ).select_related('devoir_matiere', 'devoir_matiere__matiere')

            if not reponses.exists():
                logger.warning(f"Aucune réponse corrigée pour {eleve.full_name} sur {devoir.titre}")
                return

            # Create bulletin
            bulletin = BulletinDevoir.objects.create(
                devoir=devoir,
                eleve=eleve,
                classe=classe,
                statut=BulletinDevoir.StatutBulletin.EN_ATTENTE,
                effectif_total=DevoirComposition.objects.filter(devoir=devoir).count(),
            )

            total_weighted = 0
            total_coeffs = 0

            for reponse in reponses:
                dm = reponse.devoir_matiere
                matiere_nom = dm.matiere.nom if dm.matiere else ''
                
                # Utiliser le coefficient du devoir ou le coefficient officiel béninois
                coeff_devoir = devoir.coefficients_par_matiere.get(str(dm.matiere.id) if dm.matiere else None, devoir.coefficient_default)
                coeff_officiel = get_benin_coefficient(matiere_nom, serie)
                # Utiliser le plus grand des deux
                coeff = max(float(coeff_devoir), float(coeff_officiel))
                
                note = reponse.note_finale or reponse.note_ia or 0

                appreciation = reponse.appreciation_ia or ''
                if reponse.feedback_ia and isinstance(reponse.feedback_ia, dict):
                    appreciation = reponse.feedback_ia.get('appreciation', appreciation)

                ligne = BulletinDevoirLigne.objects.create(
                    bulletin=bulletin,
                    matiere=matiere_nom,
                    coefficient=coeff,
                    note_devoir=0,  # Can be updated later
                    note_exam=note,
                    moyenne=note,
                    appreciation=appreciation[:200] if appreciation else '',
                )

                total_weighted += note * float(coeff)
                total_coeffs += float(coeff)

            # Calculate general average
            if total_coeffs > 0:
                bulletin.moyenne_generale = round(total_weighted / total_coeffs, 2)
            bulletin.decision_conseil = _get_decision(bulletin.moyenne_generale)
            bulletin.save()

            # Calculate ranks
            _calculer_rang(devoir)

            logger.info(f"Bulletin généré pour {eleve.full_name} - Moy: {bulletin.moyenne_generale}")

    except Exception as e:
        logger.error(f"Erreur génération bulletin pour {eleve.full_name}: {e}")


def _calculer_rang(devoir):
    """Calculate rank for all students in a devoir based on their bulletin averages."""
    bulletins = BulletinDevoir.objects.filter(
        devoir=devoir
    ).exclude(moyenne_generale__isnull=True).order_by('-moyenne_generale')

    for rang, bulletin in enumerate(bulletins, start=1):
        bulletin.rang = rang
        bulletin.save(update_fields=['rang'])


def _recalculer_moyenne_bulletin(bulletin):
    """Recalculate bulletin average after admin grade adjustments."""
    total_weighted = 0
    total_coeffs = 0

    for ligne in bulletin.lignes.all():
        coeff = float(ligne.coefficient)
        note = ligne.moyenne or 0
        total_weighted += note * coeff
        total_coeffs += coeff

    if total_coeffs > 0:
        bulletin.moyenne_generale = round(total_weighted / total_coeffs, 2)
    bulletin.decision_conseil = _get_decision(bulletin.moyenne_generale)
    bulletin.save(update_fields=['moyenne_generale', 'decision_conseil'])


def _generer_bulletin_pdf(bulletin):
    """Generate PDF — bulletins supprimé, no-op."""
    logger.warning("_generer_bulletin_pdf appelé mais module bulletins supprimé.")
    return

    # Code désactivé — conservé pour référence future
    eleve = bulletin.eleve
    classe = bulletin.classe or (eleve.classe if eleve else '')
    serie = classe
    
    # Appliquer les coefficients officiels aux lignes
    lignes = bulletin.lignes.all().order_by('matiere')
    for ligne in lignes:
        coeff_officiel = get_benin_coefficient(ligne.matiere, serie)
        if ligne.coefficient < coeff_officiel:
            ligne.coefficient = coeff_officiel
            ligne.save(update_fields=['coefficient'])
    
    context = {
        'bulletin': bulletin,
        'eleve': eleve,
        'devoir': bulletin.devoir,
        'classe': classe,
        'serie': serie,
        'lignes': lignes,
        'moyenne': bulletin.moyenne_generale,
        'rang': bulletin.rang,
        'effectif': bulletin.effectif_total,
        'decision': bulletin.decision_conseil,
        'mention': _get_mention(bulletin.moyenne_generale),
        'annee_scolaire': bulletin.devoir.annee_scolaire,
        'admin_phone': ADMIN_PHONE,
    }

    try:
        html = render_to_string('bulletins/bulletin_professionnel_pdf.html', context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result, link_callback=_link_callback)

        if not pdf.err:
            pdf_content = result.getvalue()
            filename = f"bulletin_{bulletin.eleve.full_name.replace(' ', '_')}_{bulletin.devoir.titre[:20]}.pdf"
            bulletin.file_pdf.save(filename, ContentFile(pdf_content), save=True)
            logger.info(f"Bulletin PDF généré pour {bulletin.eleve.full_name}")
        else:
            logger.error(f"Erreur génération PDF bulletin {bulletin.eleve.full_name}")
    except Exception as e:
        logger.error(f"Erreur génération bulletin PDF: {e}")


def _generer_certificat_devoir_pdf(cert, bulletin):
    """Generate certificate PDF with all subject grades and coefficients."""
    context = {
        'cert': cert,
        'eleve': cert.eleve,
        'devoir': cert.devoir,
        'bulletin': bulletin,
        'lignes': bulletin.lignes.all().order_by('matiere') if bulletin else [],
        'moyenne': cert.moyenne_obtenue,
        'mention': cert.mention,
        'numero': cert.numero_certificat,
        'date_delivrance': cert.date_delivrance,
        'verification_token': cert.verification_token,
        'admin_phone': ADMIN_PHONE,
        'annee_scolaire': cert.devoir.annee_scolaire if cert.devoir else timezone.now().strftime('%Y-%Y'),
    }

    try:
        html = render_to_string('devoirs/certificat_pdf.html', context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result, link_callback=_link_callback)

        if not pdf.err:
            pdf_content = result.getvalue()
            filename = f"certificat_{cert.numero_certificat}.pdf"
            cert.file_pdf.save(filename, ContentFile(pdf_content), save=True)
            logger.info(f"Certificat {cert.numero_certificat} généré.")
        else:
            logger.error(f"Erreur génération PDF certificat {cert.numero_certificat}")
    except Exception as e:
        logger.error(f"Erreur génération certificat: {e}")


def _generate_certificat_pdf(cert):
    """Legacy certificate generation — fallback."""
    bulletin = BulletinDevoir.objects.filter(
        devoir=cert.devoir, eleve=cert.eleve
    ).first()
    if bulletin:
        _generer_certificat_devoir_pdf(cert, bulletin)
        return

    context = {
        'cert': cert,
        'eleve': cert.eleve,
        'numero': cert.numero_certificat,
        'moyenne': cert.moyenne_obtenue,
        'mention': cert.mention,
        'type_certificat': cert.get_type_certificat_display(),
        'date_delivrance': cert.date_delivrance,
        'verification_token': cert.verification_token,
        'admin_phone': ADMIN_PHONE,
        'annee_scolaire': cert.devoir.annee_scolaire if cert.devoir else timezone.now().strftime('%Y-%Y'),
    }

    try:
        html = render_to_string('devoirs/certificat_pdf.html', context)
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result, link_callback=_link_callback)

        if not pdf.err:
            pdf_content = result.getvalue()
            filename = f"certificat_{cert.numero_certificat}.pdf"
            cert.file_pdf.save(filename, ContentFile(pdf_content), save=True)
            logger.info(f"Certificat {cert.numero_certificat} généré.")
        else:
            logger.error(f"Erreur génération PDF certificat {cert.numero_certificat}")
    except Exception as e:
        logger.error(f"Erreur génération certificat: {e}")
# This file should be appended to devoirs/views.py
# Copy the content below and paste it at the end of devoirs/views.py


# CLASS & SERIE MANAGEMENT VIEWS

@login_required
def classe_list_view(request):
    """Liste toutes les classes avec statistiques."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    niveau_filter = request.GET.get('niveau', '')
    classes = Classe.objects.all().order_by('niveau', 'nom')
    
    if niveau_filter:
        classes = classes.filter(niveau=niveau_filter)
    
    for classe in classes:
        classe.eleve_count = User.objects.filter(classe=classe.nom, role=User.Role.ELEVE, is_active=True).count()
    
    return render(request, 'classes/list.html', {
        'classes': classes,
        'niveau_filter': niveau_filter,
        'niveaux': ['primaire', 'secondaire', 'universitaire'],
    })


@login_required
def classe_create_view(request):
    """Crée une nouvelle classe."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    if request.method == 'POST':
        form = ClasseForm(request.POST)
        if form.is_valid():
            classe = form.save()
            messages.success(request, f"Classe '{classe.nom}' créée avec succès.")
            return redirect('classe_list')
    else:
        form = ClasseForm()
    
    return render(request, 'classes/form.html', {'form': form, 'action': 'create'})


@login_required
def classe_edit_view(request, pk):
    """Modifie une classe existante."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    classe = get_object_or_404(Classe, pk=pk)
    
    if request.method == 'POST':
        form = ClasseForm(request.POST, instance=classe)
        if form.is_valid():
            form.save()
            messages.success(request, f"Classe '{classe.nom}' modifiée avec succès.")
            return redirect('classe_list')
    else:
        form = ClasseForm(instance=classe)
    
    eleves = User.objects.filter(classe=classe.nom, role=User.Role.ELEVE, is_active=True)
    
    return render(request, 'classes/form.html', {
        'form': form, 
        'action': 'edit',
        'classe': classe,
        'eleves': eleves
    })


@login_required
def classe_delete_view(request, pk):
    """Supprime une classe (si aucun élève assigné)."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    classe = get_object_or_404(Classe, pk=pk)
    eleve_count = User.objects.filter(classe=classe.nom, role=User.Role.ELEVE, is_active=True).count()
    
    if eleve_count > 0:
        messages.error(request, f"Impossible de supprimer: {eleve_count} élève(s) assignés.")
        return redirect('classe_list')
    
    classe_nom = classe.nom
    classe.delete()
    messages.success(request, f"Classe '{classe_nom}' supprimée.")
    return redirect('classe_list')


# SCHEDULE MANAGEMENT VIEWS

@login_required
def schedule_builder_view(request, pk):
    """Interface pour programmer les horaires d'un devoir."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    devoir = get_object_or_404(Devoir, pk=pk)
    matieres = devoir.matieres.all()
    horaires = devoir.horaires or {}
    
    if request.method == 'POST':
        form = HoraireForm(request.POST)
        if form.is_valid():
            matiere = form.cleaned_data['matiere']
            date = form.cleaned_data['date'].isoformat()
            heure_demarrage = form.cleaned_data['heure_demarrage'].strftime('%H:%M')
            heure_fin = form.cleaned_data['heure_fin'].strftime('%H:%M')
            salle = form.cleaned_data.get('salle', '')
            
            for mat_nom, horaire in horaires.items():
                if horaire.get('date') == date:
                    existing_start = horaire.get('heure_demarrage', '')
                    existing_end = horaire.get('heure_fin', '')
                    if not (heure_fin <= existing_start or heure_demarrage >= existing_end):
                        messages.error(request, f"Conflit avec {mat_nom}.")
                        return redirect('schedule_builder', pk=pk)
            
            horaires[matiere.nom] = {
                'date': date,
                'heure_demarrage': heure_demarrage,
                'heure_fin': heure_fin,
                'salle': salle,
            }
            
            devoir.horaires = horaires
            devoir.save()
            messages.success(request, f"Horaire ajouté pour {matiere.nom}.")
            return redirect('schedule_builder', pk=pk)
    else:
        form = HoraireForm()
        form.fields['matiere'].queryset = matieres
    
    return render(request, 'devoirs/schedule_builder.html', {
        'devoir': devoir,
        'form': form,
        'horaires': horaires,
        'matieres_non_planifiees': matieres.exclude(nom__in=horaires.keys()),
    })


@login_required
def schedule_delete_view(request, pk, matiere_nom):
    """Supprime un créneau horaire."""
    if request.user.role != 'admin':
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    devoir = get_object_or_404(Devoir, pk=pk)
    horaires = devoir.horaires or {}
    
    if matiere_nom in horaires:
        del horaires[matiere_nom]
        devoir.horaires = horaires
        devoir.save()
        messages.success(request, f"Horaire de {matiere_nom} supprimé.")
    
    return redirect('schedule_builder', pk=pk)


# WORKFLOW DASHBOARD

@login_required
def workflow_dashboard_view(request):
    """Vue globale du workflow des devoirs."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    devoirs_total = Devoir.objects.count()
    devoirs_brouillon = Devoir.objects.filter(statut=Devoir.Statut.BROUILLON).count()
    devoirs_programmes = Devoir.objects.filter(statut=Devoir.Statut.PROGRAMME_PUBLIE).count()
    devoirs_en_cours = Devoir.objects.filter(statut=Devoir.Statut.EN_COURS).count()
    devoirs_termines = Devoir.objects.filter(statut=Devoir.Statut.TERMINE).count()
    
    epreuves_soumises = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.SOUMIS).count()
    epreuves_validees = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.VALIDE).count()
    epreuves_rejetees = DevoirMatiere.objects.filter(statut=DevoirMatiere.StatutEP.REJETE).count()
    
    copies_en_correction = DevoirReponseEleve.objects.filter(statut=DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION).count()
    copies_corrigees = DevoirReponseEleve.objects.filter(statut=DevoirReponseEleve.StatutReponse.CORRIGE).count()
    
    bulletins_en_attente = BulletinDevoir.objects.filter(statut=BulletinDevoir.StatutBulletin.EN_ATTENTE).count()
    bulletins_approuves = BulletinDevoir.objects.filter(statut=BulletinDevoir.StatutBulletin.APPROUVE).count()
    
    recent_epreuves = DevoirMatiere.objects.select_related('devoir', 'matiere', 'soumis_par').order_by('-submitted_at')[:10]
    recent_copies = DevoirReponseEleve.objects.select_related('eleve', 'devoir_matiere__devoir', 'devoir_matiere__matiere').order_by('-created_at')[:10]
    
    return render(request, 'devoirs/workflow_dashboard.html', {
        'devoirs_total': devoirs_total,
        'devoirs_brouillon': devoirs_brouillon,
        'devoirs_programmes': devoirs_programmes,
        'devoirs_en_cours': devoirs_en_cours,
        'devoirs_termines': devoirs_termines,
        'epreuves_soumises': epreuves_soumises,
        'epreuves_validees': epreuves_validees,
        'epreuves_rejetees': epreuves_rejetees,
        'copies_en_correction': copies_en_correction,
        'copies_corrigees': copies_corrigees,
        'bulletins_en_attente': bulletins_en_attente,
        'bulletins_approuves': bulletins_approuves,
        'recent_epreuves': recent_epreuves,
        'recent_copies': recent_copies,
    })


# IA LAB VIEWS

@login_required
def ia_lab_dashboard_view(request):
    """Dashboard du Laboratoire IA pour le suivi des corrections."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    copies_en_attente = DevoirReponseEleve.objects.filter(
        statut=DevoirReponseEleve.StatutReponse.SOUMIS
    ).count()
    
    copies_en_cours_list = DevoirReponseEleve.objects.filter(
        statut=DevoirReponseEleve.StatutReponse.EN_COURS_CORRECTION
    ).select_related('eleve', 'devoir_matiere__devoir', 'devoir_matiere__matiere').order_by('-created_at')[:20]
    
    copies_corrigees = DevoirReponseEleve.objects.filter(
        statut=DevoirReponseEleve.StatutReponse.CORRIGE
    ).count()
    
    copies_approuvees = DevoirReponseEleve.objects.filter(
        statut=DevoirReponseEleve.StatutReponse.APPROUVE
    ).count()
    
    copies_recentes = DevoirReponseEleve.objects.filter(
        statut__in=[DevoirReponseEleve.StatutReponse.CORRIGE, DevoirReponseEleve.StatutReponse.APPROUVE]
    ).select_related('eleve', 'devoir_matiere__devoir', 'devoir_matiere__matiere').order_by('-corrige_at')[:10]
    
    return render(request, 'ai_engine/lab_dashboard.html', {
        'copies_en_attente': copies_en_attente,
        'copies_en_cours': copies_en_cours_list,
        'copies_corrigees': copies_corrigees,
        'copies_approuvees': copies_approuvees,
        'copies_recentes': copies_recentes,
    })


@login_required
def ia_verify_copy_view(request):
    """Vérifie l'intégrité d'une copie par son ID."""
    if request.user.role not in ['admin', 'conseiller']:
        messages.error(request, "Accès refusé.")
        return redirect('dashboard')
    
    if request.method == 'POST':
        copie_id = request.POST.get('copie_id', '').strip()
        if not copie_id:
            messages.error(request, "Veuillez entrer un identifiant de copie.")
            return redirect('ia_lab_dashboard')
        
        try:
            copie = DevoirReponseEleve.objects.select_related(
                'eleve', 'devoir_matiere__devoir', 'devoir_matiere__matiere'
            ).get(copie_document_id=copie_id)
            
            messages.success(request, f"Copie trouvée: {copie.eleve.full_name} - {copie.devoir_matiere.matiere.nom}")
            
            return render(request, 'ai_engine/copy_detail.html', {
                'copie': copie,
            })
        except DevoirReponseEleve.DoesNotExist:
            messages.error(request, f"Aucune copie trouvée avec l'identifiant {copie_id}")
    
    return redirect('ia_lab_dashboard')
